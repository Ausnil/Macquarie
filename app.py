import os
from flask import Flask, render_template, request, redirect, url_for, send_file
from flask import send_from_directory, abort
from werkzeug.utils import safe_join
from werkzeug.utils import secure_filename
from datetime import datetime
import sqlite3
from utils.validation import validate_excel_file
from utils.data_processor import process_data
from utils.geolocation import add_geolocation_data
import pandas as pd
from docx import Document
# Replace this:
# from utils.data_processor import process_data

# With either:
from utils.data_processor import DataProcessor
processor = DataProcessor()

# Or keep the original import if you need backward compatibility
from utils.data_processor import process_data  # Uses the wrapper function

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['ALLOWED_EXTENSIONS'] = {'xlsx', 'xls'}
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Database configuration
DB_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'database')
os.makedirs(DB_DIR, exist_ok=True)
DB_PATH = os.path.join(DB_DIR, 'upload_logs.db')

def get_db_connection():
    """Create and return a database connection"""
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

# Initialize SQLite database for logs
def init_db():
    try:
        conn = get_db_connection()
        cursor = conn.cursor()
        cursor.execute('''
            CREATE TABLE IF NOT EXISTS upload_logs (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                timestamp TEXT NOT NULL,
                filename TEXT NOT NULL,
                customers_row_count INTEGER,
                transactions_row_count INTEGER,
                products_row_count INTEGER,
                processing_time REAL
            )
        ''')
        conn.commit()
    except sqlite3.Error as e:
        print(f"Database error: {e}")
    finally:
        if conn:
            conn.close()

init_db()

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

@app.route('/generate_report', methods=['GET', 'POST'])
def generate_report(processed_data, original_filename):
    """Generate a Word document report from processed data"""
    from docx import Document
    from docx.shared import Inches
    
    doc = Document()
    doc.add_heading('Customer Data Analysis Report', 0)
    
    # Add basic info
    doc.add_heading('File Information', level=1)
    doc.add_paragraph(f'Original filename: {original_filename}')
    doc.add_paragraph(f'Report generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
    
    # Add summary
    doc.add_heading('Key Insights', level=1)
    for insight in processed_data['summary']:
        doc.add_paragraph(insight, style='ListBullet')
    
    # Add top spenders
    doc.add_heading('Top Spenders by Category', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Category'
    hdr_cells[1].text = 'Customer Name'
    hdr_cells[2].text = 'Total Spent'
    
    for _, row in processed_data['top_spenders'].iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['category'])
        row_cells[1].text = str(row['name'])
        row_cells[2].text = f"${row['amount']:,.2f}"
    
    # Add customer rankings
    doc.add_heading('Top Customers by Total Spending', level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Rank'
    hdr_cells[1].text = 'Customer Name'
    hdr_cells[2].text = 'Total Spent'
    
    for i, (_, row) in enumerate(processed_data['customer_rankings'].head(10).iterrows(), 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = str(row['name'])
        row_cells[2].text = f"${row['total_spent']:,.2f}"
    
    # Save the document
    report_filename = f"report_{original_filename.split('.')[0]}.docx"
    report_path = os.path.join(app.config['UPLOAD_FOLDER'], report_filename)
    doc.save(report_path)
    
    return report_filename

@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'file' not in request.files:
            return redirect(request.url)
        
        file = request.files['file']
        if file.filename == '':
            return redirect(request.url)
        
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(filepath)
            
            # Validate the file
            validation_result = validate_excel_file(filepath)
            if not validation_result['valid']:
                return render_template('index.html', error=validation_result['message'])
            
            # Process start time
            start_time = datetime.now()
            
            try:
                # Read all sheets
                customers = pd.read_excel(filepath, sheet_name='Customers')
                transactions = pd.read_excel(filepath, sheet_name='Transactions')
                products = pd.read_excel(filepath, sheet_name='Products')
                
                # Process data
                processed_data = process_data(customers, transactions, products)
                
                # Add geolocation data
                processed_data['customers_with_history'] = add_geolocation_data(
                    processed_data['customers_with_history']
                )
                
                # Save processed data to new Excel
                output_filename = f"processed_{filename}"
                output_path = os.path.join(app.config['UPLOAD_FOLDER'], output_filename)
                
                with pd.ExcelWriter(output_path) as writer:
                    processed_data['customers_with_history'].to_excel(writer, sheet_name='Customers', index=False)
                    processed_data['customer_category_totals'].to_excel(writer, sheet_name='CategoryTotals', index=False)
                    processed_data['top_spenders'].to_excel(writer, sheet_name='TopSpenders', index=False)
                    processed_data['customer_rankings'].to_excel(writer, sheet_name='CustomerRankings', index=False)
                
                # Generate report
                report_filename = generate_report(processed_data, filename)
                
                # Log the upload
                end_time = datetime.now()
                processing_time = (end_time - start_time).total_seconds()
                
                try:
                    conn = get_db_connection()
                    cursor = conn.cursor()
                    cursor.execute('''
                        INSERT INTO upload_logs 
                        (timestamp, filename, customers_row_count, transactions_row_count, products_row_count, processing_time)
                        VALUES (?, ?, ?, ?, ?, ?)
                    ''', (
                        start_time.isoformat(),
                        filename,
                        len(customers),
                        len(transactions),
                        len(products),
                        processing_time
                    ))
                    conn.commit()
                except sqlite3.Error as e:
                    print(f"Failed to log upload: {e}")
                finally:
                    if conn:
                        conn.close()
                
                return render_template('results.html', 
                                     original_filename=filename,
                                     processed_filename=output_filename,
                                     report_filename=report_filename,
                                     summary=processed_data['summary'])
            
            except Exception as e:
                return render_template('index.html', error=f"Processing error: {str(e)}")
    
    return render_template('index.html')
    
@app.route('/download/<path:filename>')
def download_file(filename):
    try:
        upload_folder = os.path.abspath(app.config['UPLOAD_FOLDER'])
        app.logger.debug(f"Attempting to download from: {upload_folder}")
        app.logger.debug(f"Requested filename: {filename}")
        
        file_path = safe_join(upload_folder, filename)
        file_path = os.path.abspath(file_path)
        app.logger.debug(f"Resolved path: {file_path}")
        
        if not os.path.exists(file_path):
            app.logger.error(f"File not found: {file_path}")
            abort(404)
            
        if not os.path.isfile(file_path):
            app.logger.error(f"Path is not a file: {file_path}")
            abort(404)
            
        real_path = os.path.realpath(file_path)
        if not real_path.startswith(upload_folder):
            app.logger.error(f"Security violation: {real_path} not in {upload_folder}")
            abort(404)
            
        return send_from_directory(upload_folder, filename, as_attachment=True)
        
    except Exception as e:
        app.logger.error(f"Download failed: {str(e)}", exc_info=True)
        abort(500, "Download failed")
# ... (keep the rest of your routes and functions the same) ...

if __name__ == '__main__':
    # Ensure all directories exist
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(DB_DIR, exist_ok=True)
    app.run(debug=True)
